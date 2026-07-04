#!/usr/bin/env python3
"""
horse_test.py — 养马测试 v4.0
Hermes Agent 调教评分引擎
v4.0 = v1.0骨架 + v3.0精华（frontmatter质量分析/描述结构化评估/实际配置检测/单元测试）

核心原则：纯 stdlib，零外部依赖，只读评估零副作用。
"""

import os
import re
import json
import math
import datetime
import shutil
import subprocess
from pathlib import Path

# ============================================================
# 常量
# ============================================================
HERMES_HOME = Path.home() / ".hermes"
HISTORY_FILE = HERMES_HOME / "horse-test-history.json"
HERMES_MD = Path.home() / ".hermes.md"
SCRIPT_DIR = Path(__file__).resolve().parent

# ============================================================
# YAML Frontmatter 解析（纯 re，无 PyYAML）
# ============================================================
FRONTMATTER_RE = re.compile(r'^---\s*\n(.*?)\n(?:---|\.\.\.)\s*\n', re.DOTALL)

def _parse_simple_yaml(text):
    """用 re 解析简单 YAML key:value 和 list（无需 PyYAML）"""
    result = {}
    for line in text.split('\n'):
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        m = re.match(r'^(\w[\w_-]*)\s*:\s*(.*)', line)
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip()
            if val.startswith('"') and val.endswith('"'):
                val = val[1:-1]
            elif val.startswith("'") and val.endswith("'"):
                val = val[1:-1]
            elif val.lower() == 'true':
                val = True
            elif val.lower() == 'false':
                val = False
            elif val == '' or val == '{}':
                val = {}
            elif val.startswith('[') and val.endswith(']'):
                # 解析 YAML 行内列表 [item1, item2]
                raw = val[1:-1]
                items = []
                for item in re.split(r',\s*', raw):
                    item = item.strip().strip('"').strip("'")
                    if item: items.append(item)
                val = items
            result[key] = val
    return result if result else None

def extract_frontmatter(content):
    """提取 YAML frontmatter，返回 (dict_or_None, remaining_text)"""
    if not content or not content.startswith('---'):
        return None, content
    m = FRONTMATTER_RE.match(content)
    if not m:
        return None, content
    fm = _parse_simple_yaml(m.group(1))
    return fm, content[m.end():]

# ============================================================
# 六维评分模型
# ============================================================
DIMENSIONS = [
    {"id":"skills","icon":"🧠","name":"技能体系","weight":0.25,"items":[
        {"id":"skill_count","label":"技能数量","max":25},
        {"id":"skill_desc","label":"技能描述质量","max":25},
        {"id":"skill_coverage","label":"领域覆盖","max":20},
        {"id":"skill_org","label":"分类管理","max":15},
        {"id":"skill_community","label":"社区利用","max":15},
    ]},
    {"id":"config","icon":"⚙️","name":"配置调优","weight":0.15,"items":[
        {"id":"cfg_provider","label":"Provider配置","max":20},
        {"id":"cfg_fallback","label":"Fallback机制","max":20},
        {"id":"cfg_profile","label":"Profile管理","max":15},
        {"id":"cfg_token","label":"Token优化","max":15},
        {"id":"cfg_memory","label":"记忆管理","max":15},
        {"id":"cfg_proxy","label":"网络配置","max":15},
    ]},
    {"id":"prompt","icon":"📝","name":"Prompt工程","weight":0.15,"items":[
        {"id":"pr_sysprompt","label":"System Prompt","max":25},
        {"id":"pr_skilldesc","label":"Skill描述质量","max":25},
        {"id":"pr_context","label":"上下文管理","max":20},
        {"id":"pr_lang","label":"多语言支持","max":15},
        {"id":"pr_clarity","label":"指令精确度","max":15},
    ]},
    {"id":"tools","icon":"🔌","name":"工具生态","weight":0.20,"items":[
        {"id":"tool_plugins","label":"插件集成","max":20},
        {"id":"tool_gateway","label":"Gateway","max":15},
        {"id":"tool_cron","label":"Cron任务","max":20},
        {"id":"tool_external","label":"外部工具链","max":25},
        {"id":"tool_api","label":"API集成","max":20},
    ]},
    {"id":"operations","icon":"📊","name":"运营水平","weight":0.10,"items":[
        {"id":"op_memory","label":"记忆质量","max":25},
        {"id":"op_iteration","label":"迭代优化","max":20},
        {"id":"op_community","label":"社区参与","max":20},
        {"id":"op_docs","label":"文档化","max":20},
        {"id":"op_automation","label":"自动化程度","max":15},
    ]},
    {"id":"security","icon":"🛡️","name":"安全隐私","weight":0.15,"items":[
        {"id":"sec_privacy_rules","label":"隐私规则","max":25},
        {"id":"sec_rule_quality","label":"规则质量","max":25},
        {"id":"sec_cross_profile","label":"跨Profile保护","max":20},
        {"id":"sec_data_isolate","label":"数据隔离","max":15},
        {"id":"sec_awareness","label":"安全意识","max":15},
    ]},
]

TIERS = [
    {"min":0,"max":199,"icon":"🐣","name":"马夫","en":"Stable Boy","desc":"刚接触 Hermes，基本配置完成"},
    {"min":200,"max":349,"icon":"🐴","name":"骑手","en":"Rider","desc":"掌握基础技能，能完成日常任务"},
    {"min":350,"max":499,"icon":"🏇","name":"驯马师","en":"Horse Trainer","desc":"技能体系较完整，有多 Profile 管理"},
    {"min":500,"max":649,"icon":"🦄","name":"驯马大师","en":"Master Trainer","desc":"深度定制 Agent，熟练使用 Cron/Gateway"},
    {"min":650,"max":799,"icon":"🐉","name":"龙骑士","en":"Dragon Knight","desc":"自研技能、社区贡献、多工具链集成"},
    {"min":800,"max":1000,"icon":"👑","name":"马神","en":"Horse God","desc":"调教极致，几乎自动化的智能体工作流"},
]

# ============================================================
# 工具函数
# ============================================================
def _get_username():
    """跨平台获取用户名 (macOS: USER, Windows: USERNAME)"""
    return os.environ.get("USER") or os.environ.get("USERNAME") or "User"

def _find_process(name_pattern):
    """跨平台检测进程是否存在 (macOS: pgrep, Windows: tasklist)"""
    try:
        if os.name == 'nt':  # Windows
            r = subprocess.run(["tasklist", "/FI", f"IMAGENAME eq {name_pattern}"],
                               capture_output=True, text=True, timeout=5)
            return name_pattern.lower() in r.stdout.lower()
        else:  # macOS / Linux
            # 精确匹配：只匹配进程 cmdline，避免文件名误匹配
            r = subprocess.run(["pgrep", "-f", name_pattern],
                               capture_output=True, text=True, timeout=3)
            if r.returncode != 0: return False
            for pid in r.stdout.strip().splitlines():
                try:
                    cmd = subprocess.run(["ps", "-p", pid, "-o", "command="],
                                         capture_output=True, text=True, timeout=3).stdout
                    if re.search(name_pattern, cmd.lower()): return True
                except: continue
            return False
    except: return False

def safe_read(path, max_lines=200):
    """安全读取文件，不崩溃"""
    if not path: return ""
    try:
        with open(path,'r',encoding='utf-8') as f:
            lines = []
            for _ in range(max_lines):
                line = f.readline()
                if not line: break
                lines.append(line.rstrip())
            return '\n'.join(lines)
    except: return ""

def safe_read_all(path):
    """安全读取全部内容"""
    if not path: return ""
    try:
        with open(path,'r',encoding='utf-8') as f:
            return f.read()
    except: return ""

def analyze_frontmatter_quality(fm):
    """分析 frontmatter 字段完整度 (v4.0: 从 v3.0 采纳)"""
    if not fm or not isinstance(fm,dict):
        return 0,{}
    score = 3
    for field,pts in [("name",3),("description",3),("version",2),("author",1),("license",1),("metadata",2)]:
        if fm.get(field): score += pts
    score += min(len(fm)*2,10)
    return min(score,25),{"has_frontmatter":True,"field_count":len(fm)}

def analyze_description_quality(content):
    """分析描述文本结构化质量 (v4.0: 从 v3.0 采纳)"""
    if not content: return 0,{"length":0}
    text = content.lower()
    score = min(max(len(content)-50,0)//150+1,5)
    for kw,pts in [("usage",4),("参数",4),("example",4),("depend",3),("install",2)]:
        if kw in text: score += pts
    headers = len(re.findall(r'^#{1,3}\s+',content,re.MULTILINE))
    lists = len(re.findall(r'^\s*[-*]\s+',content,re.MULTILINE))
    code = len(re.findall(r'```',content))//2
    score += min(headers+lists+code,7)
    return min(score,25),{"length":len(content),"headers":headers,"lists":lists,"code_blocks":code}

def analyze_config_token_optimization(text):
    """分析 config.yaml 的 token 优化配置 (v4.0: 从 v3.0 采纳)"""
    has_max_tk = bool(re.search(r'max_tokens?\s*[:=]',text,re.I)) if text else False
    has_ctx = bool(re.search(r'context[_\s]?window|max_turns?\s*[:=]',text,re.I)) if text else False
    has_cmp = bool(re.search(r'compress|精简|摘要|summar',text,re.I)) if text else False
    has_dis = bool(re.search(r'disabl|禁用|排除|exclud',text,re.I)) if text else False
    score = (5 + (4 if has_max_tk else 0) + (3 if has_ctx else 0)
             + (2 if has_cmp else 0) + (1 if has_dis else 0))
    return {"has_max_tokens":has_max_tk,"has_context_window":has_ctx,
            "has_skill_compress":has_cmp,"has_tool_disable":has_dis,"token_score":min(score,15)}

def analyze_security_rules(content):
    """分析 .hermes.md 安全规则质量"""
    if not content: return {"count":0,"quality":0,"has_forbidden_paths":False,"has_sensitive_files":False,
                           "has_operation_restrict":False,"has_confirm":False}
    text = content.lower()
    count = len(re.findall(r'(?:禁止|不允许|不要|拒绝|不能|forbidden|do not)',text,re.I))
    return {"count":count,"quality":min(count*3,15),
            "has_forbidden_paths":bool(re.search(
                r'(?:禁止|不允许|不要|拒绝|不能|forbidden|do not).*(?:~/|/Users/|/home/|/var/|/etc/|路径|path|目录|dir)',
                content, re.I)),
            "has_sensitive_files":bool(re.search(r'(?:\.env|\.git|\.ssh|\.aws|secret|token|key)',text,re.I)),
            "has_operation_restrict":bool(re.search(r'\b(?:rm|del|exec|eval|sudo)\b',text,re.I)),
            "has_confirm":bool(re.search(r'确认|confirm|二次|ask',text,re.I))}

# ============================================================
# HTS 计算
# ============================================================
def calculate_hts(dim_scores, previous_score=None):
    """加权几何平均 HTS = prod(S_i^W_i) × 1000，包含趋势修正 β"""
    hts = 1.0
    for d in DIMENSIONS:
        s = max(dim_scores.get(d["id"],50)/100.0,0.001)
        hts *= math.pow(s,d["weight"])
    base = round(hts*1000,1)
    beta = 1.0
    if previous_score is not None:
        delta = base - previous_score
        if delta > 30: beta = 1.05
        elif delta < -30: beta = 0.95
    return min(round(base*beta,1),1000.0),base,beta

def get_tier(score):
    for t in TIERS:
        if t["min"] <= score <= t["max"]: return t
    return TIERS[-1]

# ============================================================
# 历史记录
# ============================================================
def load_history():
    if HISTORY_FILE.is_file():
        try: return json.loads(safe_read_all(HISTORY_FILE))
        except: return []
    return []

def save_history(hts, base, dim_scores):
    h = load_history()
    h.append({"date":datetime.datetime.now().isoformat(),"hts":hts,"base":base,"dimensions":dim_scores})
    h = h[-30:]
    HISTORY_FILE.parent.mkdir(parents=True,exist_ok=True)
    with open(HISTORY_FILE,'w',encoding='utf-8') as f:
        json.dump(h,f,ensure_ascii=False,indent=2)
    return h

# ============================================================
# 扫描器
# ============================================================
class Scanner:
    def __init__(self):
        self.data = {}
        self._scan_all()
    def _scan_all(self):
        self._skills(); self._config(); self._memory(); self._cron()
        self._plugins(); self._tools(); self._security(); self._operations()

    def _skills(self):
        d={"dirs":[],"count":0,"descriptions":[],"coverage":set(),"fm_scores":[],"dq_scores":[]}
        sk = HERMES_HOME/"skills"
        if sk.is_dir():
            for e in sorted(os.listdir(sk)):
                fp = sk/e
                if not fp.is_dir() or e.startswith('.'): continue
                d["dirs"].append(e)
                for md in ['SKILL.md','DESCRIPTION.md','README.md']:
                    mp = fp/md
                    if mp.is_file():
                        c = safe_read(mp,80)
                        d["descriptions"].append({"dir":e,"file":md,"content":c})
                        fm,_ = extract_frontmatter(c)
                        if fm:
                            sc,_ = analyze_frontmatter_quality(fm)
                            d["fm_scores"].append(sc)
                        sc,_ = analyze_description_quality(c)
                        d["dq_scores"].append(sc)
                        txt = c.lower()
                        for kw,cat in [("code","coding"),("research","research"),("web","web"),
                                       ("productivity","productivity"),("data","data-science"),
                                       ("social","social-media"),("security","security")]:
                            if kw in txt: d["coverage"].add(cat)
                        break
            d["count"]=len(d["dirs"])
        self.data["skills"]=d

    def _config(self):
        c={"provider_count":0,"has_fallback":False,"fb_ok":False,"profiles":[],"profile_count":0,
           "proxy":False,"token_opt":{},"has_cross_profile":False,"text":""}
        for p in [HERMES_HOME/"config.yaml",HERMES_HOME/"config.yml"]:
            if p.is_file():
                c["text"]=safe_read_all(p); break
        if c["text"]:
            t=c["text"]
            c["provider_count"]=len(re.findall(r'provider\s*:',t,re.I))
            c["has_fallback"]='fallback' in t.lower()
            c["fb_ok"]=not bool(re.search(r'fallback_providers\s*:\s*\[',t)) and c["has_fallback"]
            c["has_cross_profile"]=bool(re.search(r'cross[_\s]?profile',t,re.I))
            c["token_opt"]=analyze_config_token_optimization(t)
        pd = HERMES_HOME/"profiles"
        if pd.is_dir():
            c["profiles"]=[d for d in os.listdir(pd) if (pd/d).is_dir() and not d.startswith('.')]
            c["profile_count"]=len(c["profiles"])
        gc = safe_read(Path.home()/".gitconfig",100)
        c["proxy"]='proxy' in gc.lower() or '7897' in gc
        for ev in ["HTTP_PROXY","HTTPS_PROXY","http_proxy","https_proxy"]:
            if os.environ.get(ev): c["proxy"]=True; break
        self.data["config"]=c

    def _memory(self):
        m={"has_memory":False,"entries":[]}
        md = HERMES_HOME/"memory"
        if md.is_dir():
            m["entries"]=[f for f in os.listdir(md) if f.endswith(('.json','.yaml','.yml'))]
        if (HERMES_HOME/".usage.json").is_file(): m["has_memory"]=True
        self.data["memory"]=m

    def _cron(self):
        c={"tasks":[],"count":0,"has_watchdog":False,"scripts":[]}
        jp = HERMES_HOME/"cron"/"jobs.json"
        if jp.is_file():
            try:
                jd=json.loads(safe_read_all(jp))
                c["tasks"]=list(jd.keys()) if isinstance(jd,dict) else []
                c["count"]=len(c["tasks"])
                for n in c["tasks"]:
                    if 'dns' in n.lower() or 'watchdog' in n.lower(): c["has_watchdog"]=True
            except: pass
        # 跨平台脚本检测（macOS: .sh, Windows: .bat/.ps1）
        sd = HERMES_HOME / "scripts"
        script_exts = ('.sh', '.bat', '.ps1', '.cmd')
        if sd.is_dir():
            c["scripts"]=[f for f in os.listdir(sd) if f.lower().endswith(script_exts) and not f.startswith('.')]
        self.data["cron"]=c

    def _plugins(self):
        p={"installed":[],"count":0,"has_useful":False}
        pd = HERMES_HOME/"plugins"
        if pd.is_dir():
            p["installed"]=[e for e in os.listdir(pd) if (pd/e).is_dir() and not e.startswith('.')]
            p["count"]=len(p["installed"])
            p["has_useful"]=any(u in p["installed"] for u in ['disk-cleanup','security-guidance','rtk-rewrite'])
        self.data["plugins"]=p

    def _tools(self):
        t={"installed":[],"gateway":False}
        extra_paths = [str(Path.home()/".local"/"bin"), str(Path.home()/".hermes"/"node"/"bin"),
                       str(Path.home()/".hermes"/"node"), "/opt/homebrew/bin", "/usr/local/bin"]
        ex_path = ":".join(extra_paths)
        cur_path = os.environ.get("PATH","")
        for cmd in ["pipx","ffmpeg","gbrain","crush","node","bun","playwright","python3"]:
            if shutil.which(cmd, path=f"{ex_path}:{cur_path}"): t["installed"].append(cmd)
        if (Path.home()/".cache"/"ms-playwright").is_dir() and "playwright" not in t["installed"]:
            t["installed"].append("playwright")
        # Gateway 进程检测（跨平台）
        if _find_process("hermes.*gateway"): t["gateway"]=True; t["installed"].append("gateway")
        self.data["tools"]=t

    def _security(self):
        s={"has_hermes_md":HERMES_MD.is_file(),"content":"","rules_count":0,"quality":0,
           "cross_profile":False,"data_isolate":0,"has_confirm":False}
        if s["has_hermes_md"]:
            c=safe_read(HERMES_MD,150)
            s["content"]=c
            ar=analyze_security_rules(c)
            s["rules_count"]=ar["count"]
            s["quality"]=ar["quality"]
            s["has_confirm"]=ar["has_confirm"]
            iso=0
            if ar["has_forbidden_paths"]: iso+=5
            if ar["has_sensitive_files"]: iso+=5
            if ar["has_operation_restrict"]: iso+=3
            if ar["has_confirm"]: iso+=2
            s["data_isolate"]=iso
        s["cross_profile"]=self.data.get("config",{}).get("has_cross_profile",False)
        self.data["security"]=s

    def _operations(self):
        o={"history":HISTORY_FILE.is_file(),"notes":False,"community":False,"notes_count":0}
        sk=HERMES_HOME/"skills"
        if sk.is_dir():
            o["notes_count"]=len(list(sk.glob("*/使用笔记.md")))+len(list(sk.glob("*/notes.md")))
            o["notes"]=o["notes_count"]>0
            o["community"]=any(
                ('tap' in d.lower() or 'community' in d.lower() or 'github.com' in (safe_read(sk/d/"SKILL.md",20)).lower())
                for d in os.listdir(sk) if (sk/d).is_dir() and not d.startswith('.') and (sk/d/"SKILL.md").is_file()
            )
        self.data["operations"]=o

# ============================================================
# 评分器
# ============================================================
class Scorer:
    def __init__(self,scanner):
        self.d=scanner.data; self.scores={}; self.dim_scores={}
    def run(self):
        for dim in DIMENSIONS:
            fn=getattr(self,f"_s_{dim['id']}",None)
            self.scores[dim["id"]]=fn() if fn else {"score":0,"details":{i["id"]:0 for i in dim["items"]}}
            tot=sum(self.scores[dim["id"]]["details"].values())
            mx=sum(i["max"] for i in dim["items"])
            self.dim_scores[dim["id"]]=round(tot/mx*100,1) if mx else 0

    def _s_skills(self):
        s=self.d["skills"]; fm=s.get("fm_scores",[]); dq=s.get("dq_scores",[])
        det={}
        c=s["count"]
        det["skill_count"]=10 if c<=3 else 15 if c<=9 else 20 if c<=20 else 25
        if fm and dq: det["skill_desc"]=round(min(sum(fm)/len(fm)*0.4+sum(dq)/len(dq)*0.6,25))
        elif dq: det["skill_desc"]=round(min(sum(dq)/len(dq),25))
        else: det["skill_desc"]=5
        det["skill_coverage"]=min(len(s.get("coverage",set()))*4,20)
        has_desc=any(d.get("file")=="DESCRIPTION.md" for d in s.get("descriptions",[]))
        det["skill_org"]=(10 if has_desc else 5)+min(sum(1 for d in s.get("dirs",[]) if '/' in d)*3,10)
        det["skill_community"]=15 if any(
            any(kw in d.get("content","").lower() for kw in ['tap','community','wondelai','witt3rd','github'])
            for d in s.get("descriptions",[])
        ) else 8
        return {"score":sum(det.values()),"details":det}

    def _s_config(self):
        c=self.d["config"]; det={}
        det["cfg_provider"]=min(c.get("provider_count",0)*5,20)
        if c.get("has_fallback"): det["cfg_fallback"]=20 if c.get("fb_ok") else 12
        else: det["cfg_fallback"]=0
        det["cfg_profile"]=min(c.get("profile_count",0)*7,15)
        to=c.get("token_opt",{})
        det["cfg_token"]=to.get("token_score",10) if to else 10
        det["cfg_memory"]=10 if self.d.get("memory",{}).get("has_memory") else 7
        det["cfg_proxy"]=15 if c.get("proxy") else 5
        return {"score":sum(det.values()),"details":det}

    def _s_prompt(self):
        s=self.d["skills"]; fm=s.get("fm_scores",[]); dq=s.get("dq_scores",[]); det={}
        det["pr_sysprompt"]=round(min((sum(fm)/len(fm) if fm else 0),25)) if fm else 5
        det["pr_skilldesc"]=round(min((sum(dq)/len(dq) if dq else 0),25)) if dq else 10
        cfg=self.d.get("config",{}).get("text","")
        det["pr_context"]=18 if re.search(r'max_turns?\s*[:=]',cfg,re.I) else 10
        descs=s.get("descriptions",[])
        has_cn=any('中文' in d.get("content","") for d in descs)
        has_en=any(len(re.findall(r'[a-zA-Z]{8,}',d.get("content","")))>0 for d in descs)
        det["pr_lang"]=15 if has_cn and has_en else 8
        has_ex=any('example' in d.get("content","").lower() or '示例' in d.get("content","") for d in descs)
        det["pr_clarity"]=15 if has_ex else 8
        return {"score":sum(det.values()),"details":det}

    def _s_tools(self):
        det={}
        pl=self.d.get("plugins",{})
        det["tool_plugins"]=20 if pl.get("has_useful") else (10 if pl.get("count",0)>0 else 0)
        # 2. Gateway（跨平台进程检测）
        det["tool_gateway"]=15 if _find_process("hermes.*gateway") else 5
        cr=self.d.get("cron",{})
        det["tool_cron"]=min(cr.get("count",0)*7,20)
        tools=self.d.get("tools",{})
        inst=[t for t in tools.get("installed",[]) if t!="gateway"]
        det["tool_external"]=min(len(inst)*5,25)
        det["tool_api"]=min(self.d.get("config",{}).get("provider_count",0)*4,20)
        return {"score":sum(det.values()),"details":det}

    def _s_operations(self):
        det={}
        ops=self.d.get("operations",{})
        det["op_memory"]=18 if ops.get("history") else 10
        det["op_iteration"]=18 if len(load_history())>=2 else 12
        det["op_community"]=20 if ops.get("community") else 8
        det["op_docs"]=min(8+ops.get("notes_count",0)*4, 20) if ops.get("notes") else 8
        det["op_automation"]=min(self.d.get("cron",{}).get("count",0)*3,15)
        return {"score":sum(det.values()),"details":det}

    def _s_security(self):
        det={}
        sec=self.d.get("security",{})
        det["sec_privacy_rules"]=25 if sec.get("has_hermes_md") else 0
        det["sec_rule_quality"]=min(sec.get("quality",0)+sec.get("rules_count",0),25)
        det["sec_cross_profile"]=20 if sec.get("cross_profile") else 8
        det["sec_data_isolate"]=min(sec.get("data_isolate",0),15)
        det["sec_awareness"]=min(sec.get("rules_count",0)*3,15)
        return {"score":sum(det.values()),"details":det}

# ============================================================
# 建议生成
# ============================================================
def gen_suggestions(scorer):
    SUG = {
        "skills":["为每个技能补充完整 YAML frontmatter (name/description/version/author)",
                   "从 Skill Tap 安装社区高质量技能，覆盖更多领域",
                   "按目录分类组织技能，添加 DESCRIPTION.md",
                   "为技能添加 CHANGELOG.md 记录迭代"],
        "config":["在 config.yaml 中配置主备 Provider + 正确格式的 fallback_providers",
                   "创建多 Profile 分离工作/个人场景",
                   "配置 cross_profile 保护不同 Profile 的数据隔离",
                   "优化 Token：设置 max_tokens、context_window、禁用不常用工具"],
        "prompt":["为每个 SKILL.md 编写完整 frontmatter（name/description/version/metadata）",
                  "在描述中添加 Usage / Parameters / Examples 章节",
                  "配置 max_turns 参数优化上下文窗口利用",
                  "添加 Few-shot 示例提高指令精确度"],
        "tools":["安装实用插件：disk-cleanup, security-guidance, rtk-rewrite",
                 "配置 Gateway (iMessage/Telegram/Discord) 扩展触达",
                 "安装外部工具链：playwright, ffmpeg, gbrain, crush",
                 "创建 Cron 定时任务（看门狗、日报、健康检查）"],
        "operations":["优化记忆：定期清理过期条目，保持精简高信号",
                       "参与社区：分享技能到 Skill Tap，提交 PR",
                       "编写使用笔记和最佳实践记录",
                       "配置自动化运维脚本和监控任务"],
        "security":["在 ~/.hermes.md 中添加具体隐私规则（禁止路径/敏感文件/操作限制/确认机制）",
                    "在 config.yaml 中配置 cross_profile 保护",
                    "添加敏感文件隔离策略",
                    "建立危险操作二次确认机制"],
    }
    out=[]
    for dim in DIMENSIONS:
        ds=scorer.dim_scores.get(dim["id"],0)
        if ds<60:
            items=SUG.get(dim["id"],["继续优化"])
            idx=0 if ds<30 else 1 if ds<50 else 2
            out.append({"priority":"high" if ds<40 else "medium","dimension":dim["name"],
                        "icon":dim["icon"],"score":ds,"suggestion":items[min(idx,len(items)-1)]})
    out.sort(key=lambda x:0 if x["priority"]=="high" else 1)
    return out

# ============================================================
# 报告输出
# ============================================================
def print_report(scorer, hts, base, beta, tier, suggestions, history):
    user=_get_username(); now=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    print(f"\n{'='*60}")
    print(f"  🐎 养马测试结果 — Hermes Agent 调教评分报告  v4.0")
    print(f"{'='*60}")
    print(f"  用户: {user}        日期: {now}")
    print(f"  HTS 综合评分: {hts}/1000  |  等级: {tier['icon']} {tier['name']} ({tier['en']})")
    print(f"{'='*60}\n")
    print(f"  📊 六维评分详情:")
    print(f"  {'─'*58}")
    for dim in DIMENSIONS:
        ds=scorer.dim_scores.get(dim["id"],0)
        bar="█"*int(ds/100*30)+"░"*(30-int(ds/100*30))
        print(f"  {dim['icon']} {dim['name']:<12s} {bar} {ds:5.1f}/100 (权重{dim['weight']*100:.0f}%)")
    print(f"  {'─'*58}")
    print(f"  📐 算法: 加权几何平均 HTS = ∏(Sᵢ^Wᵢ)")
    print(f"     基础 HTS: {base}  |  趋势修正 β: {beta:.2f}  |  最终得分: {hts}")
    if history:
        p_base=history[-1].get("base",0); p_hts=history[-1].get("hts",0); d=base-p_base
        print(f"     上次评分: HTS {p_hts}  |  基础分: {p_base} → {base} {'📈' if d>0 else '📉' if d<0 else '➖'} {d:+.1f}")
    print(f"\n  🏆 当前段位: {tier['icon']} {tier['name']} (第{TIERS.index(tier)+1}段/共{len(TIERS)}段)")
    print(f"     {tier['desc']}")
    nxt=TIERS.index(tier)+1
    if nxt<len(TIERS):
        need=TIERS[nxt]["min"]-hts; print(f"     距离下一段「{TIERS[nxt]['icon']}{TIERS[nxt]['name']}」还需 {need:.1f} 分")
    if suggestions:
        print(f"\n  💡 改进建议:")
        for i,s in enumerate(suggestions[:5],1):
            tag="🔴" if s["priority"]=="high" else "🟡"
            print(f"  {i}. {tag} {s['icon']} {s['dimension']} ({s['score']:.0f}/100)")
            print(f"     → {s['suggestion']}")
    print(f"\n{'='*60}")

def gen_word_report(scorer, hts, base, beta, tier, suggestions, dim_scores, sdata):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  ⚠️ python-docx 未安装，跳过 Word 报告")
        return None
    doc=Document()
    def sf(run,sz=Pt(10),color=None,bold=False):
        run.font.name='Arial'; run.font.size=sz; run.bold=bold
        if color: run.font.color.rgb=color
        r=run._element; rPr=r.get_or_add_rPr(); rFonts=rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'),'Arial')
    def sch(cell,color):
        s=OxmlElement('w:shd'); s.set(qn('w:fill'),color); s.set(qn('w:val'),'clear')
        cell._element.get_or_add_tcPr().append(s)
    def hdr(tbl,texts):
        for i,t in enumerate(texts):
            tbl.rows[0].cells[i].text=t; sch(tbl.rows[0].cells[i],'1A3C6E')
            for p in tbl.rows[0].cells[i].paragraphs:
                for r in p.runs: sf(r,Pt(10),RGBColor(0xFF,0xFF,0xFF),True)
    # 封面
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('🐎 养马测试报告 v4.0'); sf(r,Pt(30),RGBColor(0x8B,0x45,0x13),True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'HTS: {hts}/1000  |  {tier["icon"]} {tier["name"]} ({tier["en"]})'); sf(r,Pt(16))
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'生成: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}'); sf(r,Pt(10),RGBColor(0x66,0x66,0x66))
    doc.add_page_break()
    # 评分总览
    doc.add_heading('📊 评分总览',level=1)
    tbl=doc.add_table(rows=7,cols=4); tbl.style='Light Grid Accent 1'
    hdr(tbl,['维度','得分','权重','评价'])
    for ri,dim in enumerate(DIMENSIONS):
        ds=dim_scores.get(dim["id"],0); tbl.rows[ri+1].cells[0].text=f'{dim["icon"]} {dim["name"]}'
        tbl.rows[ri+1].cells[1].text=f'{ds:.1f}/100'; tbl.rows[ri+1].cells[2].text=f'{dim["weight"]*100:.0f}%'
        tbl.rows[ri+1].cells[3].text='⭐ 优秀' if ds>=80 else '✅ 良好' if ds>=60 else '🔧 一般' if ds>=40 else '🔴 待改进'
    tbl.rows[6].cells[0].text='🏆 综合(HTS)'; tbl.rows[6].cells[1].text=f'{hts}/1000'
    tbl.rows[6].cells[2].text='100%'; tbl.rows[6].cells[3].text=f'{tier["icon"]} {tier["name"]}'
    doc.add_paragraph(f'算法: 加权几何平均 | 基础HTS: {base} | β: {beta:.2f}')
    doc.add_page_break()
    # 评分明细
    doc.add_heading('📋 各维度评分明细',level=1)
    for dim in DIMENSIONS:
        did=dim["id"]; ds=dim_scores.get(did,0)
        doc.add_heading(f'{dim["icon"]} {dim["name"]} — {ds:.1f}/100 (权重{dim["weight"]*100:.0f}%)',level=2)
        dets=scorer.scores.get(did,{}).get("details",{})
        if not dets:
            doc.add_paragraph(f'⚠️ 维度 {dim["name"]}({did}) 评分数据缺失，请检查 Scanner 扫描结果')
            continue
        items=dim["items"]
        t=doc.add_table(rows=len(items)+1,cols=4); t.style='Light Grid Accent 1'
        hdr(t,['评估项','得分','满分','评分标准'])
        for ri,item in enumerate(items):
            sc=dets.get(item["id"],0)
            t.rows[ri+1].cells[0].text=item["label"]
            t.rows[ri+1].cells[1].text=str(sc)
            t.rows[ri+1].cells[2].text=str(item["max"])
            # 找对应的评分标准（从 DIMENSIONS 的描述里来）
            criteria_map={
                "skill_count":"0-3:10, 4-9:15, 10-20:20, 20+:25",
                "skill_desc":"frontmatter完整度+描述结构化分析",
                "skill_coverage":"领域覆盖数×4",
                "skill_org":"DESCRIPTION.md(10)+子目录(3)+CHANGELOG(3)",
                "skill_community":"检测community/tap/github关键词",
                "cfg_provider":"Provider数×5",
                "cfg_fallback":"有且格式正确(20)/有但格式错(12)/无(0)",
                "cfg_profile":"Profile数×7",
                "cfg_token":"实际检测max_tokens/context/compress/disable",
                "cfg_memory":"有记忆文件(10)/无(7)",
                "cfg_proxy":"有代理配置(15)/无(5)",
                "pr_sysprompt":"frontmatter平均分(0-25)",
                "pr_skilldesc":"描述质量平均分(0-25)",
                "pr_context":"有max_turns(18)/无(10)",
                "pr_lang":"双语(15)/单语(8)",
                "pr_clarity":"有示例(15)/无(8)",
                "tool_plugins":"实用插件(20)/有插件(10)/无(0)",
                "tool_gateway":"运行中(15)/未运行(5)",
                "tool_cron":"Cron数×7",
                "tool_external":"外部工具数×5",
                "tool_api":"Provider数×4",
                "op_memory":"有历史(18)/无(10)",
                "op_iteration":"≥2次(18)/单次(12)",
                "op_community":"有(20)/无(8)",
                "op_docs":"有笔记(20+)/无(8)",
                "op_automation":"Cron数×3",
                "sec_privacy_rules":"有.hermes.md(25)/无(0)",
                "sec_rule_quality":"quality+rules_count(0-25)",
                "sec_cross_profile":"已配置(20)/未配置(8)",
                "sec_data_isolate":".hermes.md四维分析(0-15)",
                "sec_awareness":"规则数×3(0-15)",
            }
            t.rows[ri+1].cells[3].text=criteria_map.get(item["id"],item.get("criteria",""))
        doc.add_paragraph()
    # 建议
    doc.add_heading('💡 改进建议',level=1)
    if suggestions:
        for i,s in enumerate(suggestions[:8],1):
            tag='🔴' if s["priority"]=="high" else '🟡'
            p=doc.add_paragraph(); r=p.add_run(f'{i}. {tag} {s["icon"]} {s["dimension"]}'); sf(r,bold=True)
            doc.add_paragraph(f'   → {s["suggestion"]}')
    else: doc.add_paragraph('✅ 各维度表现优秀！')
    desktop = Path.home() / "Desktop"
    if not desktop.is_dir():
        desktop = Path.home() / "桌面"
        if not desktop.is_dir():
            desktop = Path.home()  # fallback to home
    path=desktop/f'养马测试报告_v4_{datetime.date.today().strftime("%Y%m%d")}.docx'
    doc.save(path); return str(path)

# ============================================================
# 主入口
# ============================================================
def main():
    import argparse
    ap = argparse.ArgumentParser(description="养马测试 v4.0 — Hermes Agent 调教评分系统")
    ap.add_argument("--json", action="store_true", help="仅输出 JSON 评分结果")
    ap.add_argument("--no-word", action="store_true", help="跳过 Word 报告生成")
    ap.add_argument("--quiet", action="store_true", help="仅输出总分和等级")
    ap.add_argument("--hermes-home", type=str, default=None,
                    help="自定义 Hermes 主目录路径（默认：~/.hermes），其他 AI agent 可用此参数指定路径")
    args = ap.parse_args()

    # 支持自定义 HERMES_HOME（其他 AI agent 可用 --hermes-home 指定路径）
    if args.hermes_home:
        global HERMES_HOME, HISTORY_FILE, HERMES_MD
        HERMES_HOME = Path(args.hermes_home).resolve()
        HISTORY_FILE = HERMES_HOME / "horse-test-history.json"
        HERMES_MD = HERMES_HOME.with_name(".hermes.md")

    if not args.quiet:
        print("🔍 养马测试 v4.0 — 正在扫描 Hermes Agent 环境...")
    sc=Scanner(); sr=Scorer(sc); sr.run()
    h=load_history(); prev_base=h[-1].get("base") if h else None; prev_hts=h[-1].get("hts") if h else None
    # 旧格式历史无 base 字段 → 不使用趋势修正（防止 base=0 触发异常 delta）
    if prev_base is not None and prev_base == 0 and h and "base" not in h[-1]:
        prev_base = None
    hts,base,beta=calculate_hts(sr.dim_scores,prev_base)
    tier=get_tier(hts); sug=gen_suggestions(sr) if base<800 else []

    if args.json:
        import json as j
        print(j.dumps({"hts":hts,"base":base,"beta":beta,
            "tier":{"name":tier["name"],"en":tier["en"],"icon":tier["icon"]},
            "dimensions":{d["id"]:sr.dim_scores.get(d["id"],0) for d in DIMENSIONS}},
            ensure_ascii=False,indent=2))
    elif args.quiet:
        print(f"HTS={hts}|TIER={tier['name']}|{tier['icon']}")
    else:
        print_report(sr,hts,base,beta,tier,sug,h)
    save_history(hts,base,sr.dim_scores)
    if not args.no_word:
        wp=gen_word_report(sr,hts,base,beta,tier,sug,sr.dim_scores,sc.data)
        if wp: print(f"\n  📄 Word 报告: {wp}")
    if not args.quiet and not args.json:
        print(f"\n💬 SUMMARY|HTS={hts}|TIER={tier['name']}|DIMS=",end="")
        for dim in DIMENSIONS: print(f"{dim['id']}={sr.dim_scores.get(dim['id'],0):.0f}",end=",")
        print()

if __name__=="__main__":
    main()
